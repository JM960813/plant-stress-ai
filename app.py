import io

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
import streamlit as st
from matplotlib.lines import Line2D
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook

plt.rcParams.update(
    {
        "figure.facecolor": "#f3faf3",
        "axes.facecolor": "#f3faf3",
        "savefig.facecolor": "#f3faf3",
        "text.color": "black",
        "axes.labelcolor": "black",
        "axes.edgecolor": "black",
        "axes.titlecolor": "black",
        "xtick.color": "black",
        "ytick.color": "black",
        "grid.color": "gray",
        "grid.alpha": 0.2,
    }
)

# =========================
# APP CONFIG
# =========================
st.set_page_config(page_title="Plant Stress AI", layout="wide")

st.markdown(
    """
<style>

/* Fondo general */
.stApp {
    background-color: #f3faf3;
}

/* Contenedores */
.block-container {
    padding-top: 2rem;
    padding-bottom: 2rem;
}

div[data-testid="stDataFrame"] {
    background-color: #ffffff;
    border: 1px solid #d8e6d8;
    border-radius: 14px;
    padding: 0.5rem;
    box-shadow: 0 8px 20px rgba(46, 125, 50, 0.06);
}

div[data-testid="stMarkdownContainer"] table {
    border-collapse: separate;
    border-spacing: 0;
}

div[data-testid="stDataFrame"] thead th,
div[data-testid="stDataFrame"] [role="columnheader"] {
    background-color: #2E7D32 !important;
    color: white !important;
    font-weight: 700 !important;
    border-color: #2E7D32 !important;
}

div[data-testid="stDataFrame"] [role="gridcell"],
div[data-testid="stDataFrame"] tbody td {
    color: #1b1b1b !important;
}

/* Texto general */
html, body, [class*="css"] {
    color: #1b1b1b;
}

/* Botones */
.stButton>button {
    background-color: #2E7D32;
    color: white;
    border-radius: 8px;
    border: none;
}

.stButton>button:hover {
    background-color: #1B5E20;
}

/* Títulos */
h1, h2, h3 {
    color: #1f3d1f;
}

.hero-card {
    background:
        radial-gradient(circle at top right, rgba(168, 230, 163, 0.55), transparent 28%),
        linear-gradient(135deg, #ffffff 0%, #eef7ee 52%, #e3f2e3 100%);
    border: 1px solid #d8e6d8;
    border-radius: 28px;
    padding: 1.7rem 1.8rem 1.9rem 1.8rem;
    box-shadow: 0 22px 48px rgba(46, 125, 50, 0.11);
    margin-bottom: 1.15rem;
    overflow: hidden;
    position: relative;
}

.hero-card::after {
    content: "";
    position: absolute;
    left: 0;
    right: 0;
    bottom: 0;
    height: 18px;
    background:
        linear-gradient(90deg,
            #2E7D32 0%,
            #3c9140 18%,
            #5aa65d 36%,
            #7fbe6f 54%,
            #a8d58d 72%,
            #d4e9bb 100%);
    opacity: 0.95;
}

.hero-badge {
    display: inline-block;
    font-size: 0.84rem;
    font-weight: 700;
    letter-spacing: 0.04em;
    text-transform: uppercase;
    color: #1f5c2b;
    background: rgba(46, 125, 50, 0.10);
    border: 1px solid rgba(46, 125, 50, 0.18);
    border-radius: 999px;
    padding: 0.35rem 0.8rem;
    margin-bottom: 0.85rem;
}

.hero-layout {
    display: grid;
    grid-template-columns: minmax(0, 1.55fr) minmax(300px, 0.85fr);
    gap: 1.25rem;
    align-items: stretch;
}

.hero-copy {
    position: relative;
    z-index: 1;
}

.hero-title {
    font-size: 2rem;
    font-weight: 800;
    line-height: 1.08;
    color: #17351d;
    margin: 0 0 0.45rem 0;
}

.hero-subtitle {
    font-size: 1rem;
    line-height: 1.55;
    color: #35523a;
    max-width: 900px;
    margin: 0 0 1rem 0;
}

.hero-pills {
    display: flex;
    flex-wrap: wrap;
    gap: 0.65rem;
}

.hero-pill {
    background: #ffffff;
    border: 1px solid #d8e6d8;
    border-radius: 999px;
    padding: 0.5rem 0.85rem;
    color: #24442a;
    font-size: 0.92rem;
    font-weight: 600;
}

.hero-visual {
    position: relative;
    min-height: 250px;
    border-radius: 24px;
    border: 1px solid rgba(46, 125, 50, 0.14);
    background:
        radial-gradient(circle at 70% 20%, rgba(255,255,255,0.85), transparent 26%),
        radial-gradient(circle at 35% 28%, rgba(196, 224, 163, 0.82), transparent 44%),
        linear-gradient(180deg, #dceecf 0%, #d7e7bd 34%, #6d5a43 35%, #4d3b2c 100%);
    box-shadow: inset 0 0 0 1px rgba(255,255,255,0.22);
    overflow: hidden;
}

.hero-visual::before {
    content: "";
    position: absolute;
    inset: auto 0 0 0;
    height: 42%;
    background:
        radial-gradient(circle at 18% 0%, rgba(130, 92, 62, 0.42), transparent 22%),
        radial-gradient(circle at 52% 10%, rgba(65, 44, 31, 0.34), transparent 20%),
        radial-gradient(circle at 78% 0%, rgba(126, 88, 58, 0.30), transparent 22%);
}

.hero-plant {
    position: absolute;
    left: 50%;
    top: 52%;
    transform: translate(-50%, -50%);
    font-size: 6.1rem;
    filter: drop-shadow(0 8px 12px rgba(31, 61, 31, 0.18));
}

.hero-findings {
    position: absolute;
    right: 1rem;
    top: 1rem;
    width: min(86%, 340px);
    background: rgba(255, 255, 255, 0.93);
    border: 1px solid rgba(46, 125, 50, 0.16);
    border-radius: 22px;
    padding: 1rem 1rem 0.95rem 1rem;
    box-shadow: 0 16px 32px rgba(46, 125, 50, 0.10);
    backdrop-filter: blur(3px);
}

.hero-findings-title {
    font-size: 0.88rem;
    font-weight: 800;
    letter-spacing: 0.04em;
    text-transform: uppercase;
    color: #1f5c2b;
    margin-bottom: 0.8rem;
}

.hero-findings-grid {
    display: grid;
    grid-template-columns: repeat(3, minmax(0, 1fr));
    gap: 0.55rem;
}

.hero-finding {
    text-align: center;
    padding: 0.35rem 0.3rem 0.15rem 0.3rem;
}

.hero-finding-icon {
    font-size: 1.55rem;
    display: block;
    margin-bottom: 0.28rem;
}

.hero-finding-label {
    font-size: 0.72rem;
    font-weight: 700;
    color: #47614c;
    line-height: 1.25;
}

.hero-finding-value {
    font-size: 1.18rem;
    font-weight: 800;
    color: #17351d;
    margin-top: 0.25rem;
}

.section-card {
    background: rgba(255, 255, 255, 0.78);
    border: 1px solid #d8e6d8;
    border-radius: 20px;
    padding: 1.05rem 1.1rem 1.2rem 1.1rem;
    box-shadow: 0 10px 24px rgba(46, 125, 50, 0.05);
    margin: 0.9rem 0 1.15rem 0;
}

.section-kicker {
    font-size: 0.78rem;
    font-weight: 800;
    letter-spacing: 0.05em;
    text-transform: uppercase;
    color: #2d6a36;
    margin-bottom: 0.45rem;
}

.metrics-card {
    background: linear-gradient(180deg, #ffffff 0%, #f7fbf7 100%);
    border: 1px solid #d8e6d8;
    border-radius: 24px;
    padding: 1rem 1.05rem 1.1rem 1.05rem;
    box-shadow: 0 16px 30px rgba(46, 125, 50, 0.07);
}

.metrics-note {
    color: #47614c;
    font-size: 0.9rem;
    margin: 0.25rem 0 0.8rem 0;
}

.summary-grid {
    display: grid;
    grid-template-columns: repeat(4, minmax(0, 1fr));
    gap: 0.85rem;
}

.summary-tile {
    background: linear-gradient(180deg, #ffffff 0%, #fcfefb 100%);
    border: 1px solid #dfeadf;
    border-radius: 18px;
    padding: 0.95rem 0.9rem;
    min-height: 120px;
    box-shadow: inset 0 1px 0 rgba(255,255,255,0.7);
}

.summary-kicker {
    font-size: 0.78rem;
    font-weight: 800;
    text-transform: uppercase;
    letter-spacing: 0.04em;
    color: #55705a;
    margin-bottom: 0.45rem;
}

.summary-value {
    font-size: 1.55rem;
    font-weight: 800;
    color: #17351d;
    line-height: 1.1;
}

.summary-detail {
    margin-top: 0.35rem;
    color: #48624d;
    font-size: 0.88rem;
    line-height: 1.35;
}

.spotlight-card {
    background: linear-gradient(180deg, #ffffff 0%, #f8fcf7 100%);
    border: 1px solid #d8e6d8;
    border-radius: 24px;
    padding: 1rem 1.05rem 1.1rem 1.05rem;
    box-shadow: 0 16px 30px rgba(46, 125, 50, 0.07);
}

.spotlight-header {
    font-size: 0.82rem;
    font-weight: 800;
    letter-spacing: 0.04em;
    text-transform: uppercase;
    color: #2a5a33;
    margin-bottom: 0.4rem;
}

.spotlight-value {
    font-size: 1.75rem;
    font-weight: 800;
    color: #17351d;
    margin-bottom: 0.25rem;
}

.spotlight-note {
    color: #4b6650;
    font-size: 0.9rem;
    line-height: 1.45;
    margin-bottom: 0.85rem;
}

.feature-strip {
    display: grid;
    grid-template-columns: repeat(4, minmax(0, 1fr));
    gap: 0.9rem;
    margin-top: 1rem;
}

.feature-card {
    background: rgba(255,255,255,0.72);
    border: 1px solid #d8e6d8;
    border-radius: 18px;
    padding: 0.95rem 1rem;
    box-shadow: 0 10px 22px rgba(46, 125, 50, 0.05);
}

.feature-card strong {
    display: block;
    color: #17351d;
    margin-bottom: 0.22rem;
}

.feature-card span {
    color: #4a6250;
    font-size: 0.9rem;
}

.leader-grid {
    display: grid;
    grid-template-columns: repeat(3, minmax(0, 1fr));
    gap: 0.85rem;
    margin: 0.35rem 0 1rem 0;
}

.leader-card {
    background: linear-gradient(180deg, #ffffff 0%, #f9fcf8 100%);
    border: 1px solid #dce8dc;
    border-radius: 18px;
    padding: 0.95rem 1rem;
    box-shadow: 0 10px 20px rgba(46, 125, 50, 0.05);
}

.leader-rank {
    font-size: 0.78rem;
    font-weight: 800;
    letter-spacing: 0.04em;
    text-transform: uppercase;
    color: #56715b;
    margin-bottom: 0.3rem;
}

.leader-genotype {
    font-size: 1.35rem;
    font-weight: 800;
    color: #17351d;
}

.leader-score {
    margin-top: 0.35rem;
    color: #49644e;
    font-size: 0.92rem;
}

.executive-card {
    background: linear-gradient(180deg, #ffffff 0%, #f8fcf7 100%);
    border: 1px solid #d8e6d8;
    border-radius: 20px;
    padding: 1rem 1.05rem;
    box-shadow: 0 12px 24px rgba(46, 125, 50, 0.06);
    margin: 0.7rem 0 0.85rem 0;
}

.executive-title {
    font-size: 0.84rem;
    font-weight: 800;
    text-transform: uppercase;
    letter-spacing: 0.05em;
    color: #2a5a33;
    margin-bottom: 0.4rem;
}

.executive-text {
    color: #345039;
    line-height: 1.55;
    font-size: 0.95rem;
}

@media (max-width: 900px) {
    .hero-layout,
    .summary-grid,
    .feature-strip,
    .leader-grid {
        grid-template-columns: 1fr;
    }
    .hero-findings-grid {
        grid-template-columns: 1fr;
    }
    .hero-findings {
        position: static;
        width: auto;
        margin: 1rem;
    }
    .hero-visual {
        min-height: 340px;
    }
}

</style>
""",
    unsafe_allow_html=True,
)

st.markdown(
    """
<div class="hero-card">
  <div class="hero-layout">
    <div class="hero-copy">
      <div class="hero-badge">🌾 Plant Phenotyping Dashboard</div>
      <div class="hero-title">🌱 Drought Stress Phenotyping &amp; Breeding AI</div>
      <div class="hero-subtitle">
        A crop-oriented scientific dashboard for identifying drought-tolerant genotypes
        using <b>Chlorophyll Content</b>, <b>Photosystem II Efficiency</b>, <b>Yield</b>,
        and integrated stress-response modeling.
      </div>
      <div class="hero-pills">
        <div class="hero-pill">🍃 Chlorophyll Content</div>
        <div class="hero-pill">🌿 Photosystem II Efficiency</div>
        <div class="hero-pill">🌾 Yield</div>
        <div class="hero-pill">🧪 Stress Index</div>
        <div class="hero-pill">🎯 Breeding Decision Support</div>
      </div>
    </div>
    <div class="hero-visual">
      <div class="hero-plant">🌱</div>
      <div class="hero-findings">
        <div class="hero-findings-title">Scientific Workflow</div>
        <div class="hero-findings-grid">
          <div class="hero-finding">
            <span class="hero-finding-icon">🧬</span>
            <div class="hero-finding-label">Trait<br>screening</div>
            <div class="hero-finding-value">1</div>
          </div>
          <div class="hero-finding">
            <span class="hero-finding-icon">📊</span>
            <div class="hero-finding-label">Stress + Yield<br>analysis</div>
            <div class="hero-finding-value">2</div>
          </div>
          <div class="hero-finding">
            <span class="hero-finding-icon">🥇</span>
            <div class="hero-finding-label">Breeding<br>decision</div>
            <div class="hero-finding-value">3</div>
          </div>
        </div>
      </div>
    </div>
  </div>
</div>
""",
    unsafe_allow_html=True,
)

st.markdown(
    """
<style>
/* ajustes estables */
.block-container {
    background-color: transparent;
}

.main {
    background-color: transparent;
}

/* texto más legible en ambos modes */
p, label, div {
    color: inherit;
}
</style>
""",
    unsafe_allow_html=True,
)

st.info(
    """
This tool integrates Chlorophyll Content, Photosystem II Efficiency, and Yield with stress modeling to support drought-tolerant genotype selection.

It is crop-independent and can be applied to any plant species as long as the required physiological and yield data are provided.
"""
)

with st.expander("📌 How to use this app"):
    st.markdown(
        """
1. Download the Excel template
2. Fill in your experimental data:
   - Genotype (e.g. H1, H2...)
   - Chlorophyll_Content
   - Photosystem_II_Efficiency
   - Yield (recommended for breeding and productivity analysis)
3. Upload the file
4. Click Run Analysis
5. View results and download report
"""
    )


with st.expander("🧠 Interpretation Guide"):
    st.markdown(
        """
### Updated Trait Definitions

- **Chlorophyll Content (formerly SPAD):**  
  Measures leaf chlorophyll concentration and greenness. Higher values = healthier plants.

- **Photosystem II Efficiency (formerly Fv/Fm):**  
  Measures photosynthetic efficiency under stress conditions. Higher values = better physiological performance.

- **Yield:**  
  Represents productivity under the evaluated conditions. Higher values = stronger agronomic performance.

👉 These variables are complementary: the physiological traits describe plant health, while Yield helps confirm whether that tolerance is translated into productivity.
"""
    )

with st.expander("🌿 What do variables mean?"):
    st.markdown(
        """
### Photosystem II Efficiency
- Measures how efficiently the plant performs photosynthesis  
- Higher values = healthier plant  
- Lower values = stronger physiological stress damage  

### Chlorophyll Content
- Estimates chlorophyll level and leaf greenness  
- Higher values = greener, healthier leaves  
- Lower values = stronger stress or reduced physiological performance  

### Stress Index
The app combines both normalized physiological variables:

**Stress Index = 1 - (normalized Chlorophyll Content + normalized Photosystem II Efficiency) / 2**

- Close to 0 → tolerant genotype  
- Close to 1 → sensitive genotype  

### Yield
- Represents the productive outcome of each genotype under the evaluated conditions  
- Higher values = better agronomic performance  
- When combined with Stress Index, it helps identify elite genotypes with both tolerance and productivity  
"""
    )

st.success(
    """
The model automatically classifies genotypes into Low, Moderate, and High stress groups based on Stress Index distribution.
"""
)


# =========================
# TEMPLATE DOWNLOAD
# =========================
st.subheader("📥 Step 1: Download Template")

wb = Workbook()
ws = wb.active
ws.title = "template"

ws.append(
    [
        "Genotype",
        "Chlorophyll_Content",
        "Photosystem_II_Efficiency",
        "Yield",
    ]
)

for i in range(1, 6):
    ws.append([f"H{i}", "", "", ""])

buffer = io.BytesIO()
wb.save(buffer)
buffer.seek(0)

st.download_button(
    label="📥 Download Template",
    data=buffer,
    file_name="phenotyping_template.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
)


# =========================
# CORE ANALYSIS
# =========================
SPAD_COL = "Chlorophyll_Content"
FVFM_COL = "Photosystem_II_Efficiency"
LOW_STRESS_COLOR = "#2E7D32"
MODERATE_STRESS_COLOR = "#F9A825"
HIGH_STRESS_COLOR = "#C62828"
DASHBOARD_BG = "#f3faf3"
TEXT_COLOR = "#1b1b1b"
BORDER_COLOR = "#d8e6d8"
CARD_BG = "#ffffff"
YIELD_COLOR = "#6FAF73"


def normalize_trait_columns(df):
    df = df.copy()
    df.columns = df.columns.str.strip()
    return df.rename(
        columns={
            "SPAD": SPAD_COL,
            "FvFm": FVFM_COL,
        }
    )


def compute_stress(df):
    df = df.copy()

    df["FvFm_norm"] = df[FVFM_COL] / df[FVFM_COL].max()
    df["SPAD_norm"] = df[SPAD_COL] / df[SPAD_COL].max()
    df["Stress_Index"] = 1 - (df["FvFm_norm"] + df["SPAD_norm"]) / 2

    return df


def classify(x):
    if x < 0.4:
        return "🟢 Low"
    if x < 0.7:
        return "🟡 Moderate"
    return "🔴 High"


def generate_recommendation(avg):
    if avg < 0.4:
        return "Best group: high drought tolerance"
    if avg < 0.7:
        return "Moderate stress response"
    return "High stress sensitivity detected"


def color_scale(val):
    if val < 0.4:
        return LOW_STRESS_COLOR
    if val < 0.7:
        return MODERATE_STRESS_COLOR
    return HIGH_STRESS_COLOR


def stress_class_name(val):
    if val < 0.4:
        return "Low stress"
    if val < 0.7:
        return "Moderate stress"
    return "High stress"


def color_class(val):
    return f"background-color: {color_scale(val)}"


def safe_yield_available(df):
    return (
        "Yield" in df.columns
        and df["Yield"].notna().any()
        and df["Yield"].sum() > 0
    )


def run_yield_model(df):
    if not safe_yield_available(df):
        return None


def get_color(si, y, df):
    if si < df["Stress_Index"].quantile(0.33) and y > df["Yield"].quantile(0.66):
        return LOW_STRESS_COLOR
    if si > df["Stress_Index"].quantile(0.66) and y < df["Yield"].quantile(0.33):
        return HIGH_STRESS_COLOR
    return MODERATE_STRESS_COLOR


def explain_stress_vs_yield():
    st.markdown(
        """
    ### ⚖️ Stress vs Yield Trade-off (Breeding Decision Map)

    ### 🧠 How to interpret this figure

    This graph combines drought stress and productivity in one view.

    - **X-axis:** Stress Index (higher = more sensitive genotypes)
    - **Y-axis:** Yield (higher = more productive genotypes)

    ### 📌 Breeding interpretation:
    - 🟢 Upper-left → ideal genotypes (low stress + high yield)
    - 🔴 Lower-right → poor performers (high stress + low yield)
    - 🟡 Middle → intermediate performance

    👉 This plot is used to select **elite breeding candidates**.
    """
    )


def trait_texts():
    st.markdown(
        """
### 🌿 Physiological Traits (Standardized Definitions)

- **Chlorophyll Content (SPAD equivalent):**  
  Indicates leaf chlorophyll concentration and greenness.

- **Photosystem II Efficiency (Fv/Fm):**  
  Measures photosynthetic performance under stress conditions.

- **Stress Index:**  
  Quantifies reduction in plant performance under drought stress.
"""
    )


def style_matplotlib(ax):
    import matplotlib.pyplot as plt

    fig = plt.gcf()

    fig.patch.set_facecolor("#f3faf3")
    ax.set_facecolor("#f3faf3")

    ax.tick_params(colors="black")
    ax.title.set_color("black")
    ax.xaxis.label.set_color("black")
    ax.yaxis.label.set_color("black")
    for spine in ax.spines.values():
        spine.set_color("black")

    ax.grid(True, color="gray", alpha=0.2)


def style_colorbar(colorbar):
    colorbar.ax.tick_params(colors="black")
    colorbar.ax.yaxis.label.set_color("black")
    colorbar.outline.set_edgecolor("black")


def style_dataframe(styler):
    return styler.set_table_styles(
        [
            {
                "selector": "thead th",
                "props": [
                    ("background-color", LOW_STRESS_COLOR),
                    ("color", "white"),
                    ("font-weight", "700"),
                    ("border", f"1px solid {LOW_STRESS_COLOR}"),
                    ("padding", "11px 12px"),
                    ("text-align", "left"),
                    ("font-size", "0.9rem"),
                    ("letter-spacing", "0.01em"),
                ],
            },
            {
                "selector": "th.blank",
                "props": [
                    ("background-color", LOW_STRESS_COLOR),
                    ("color", "white"),
                    ("border", f"1px solid {LOW_STRESS_COLOR}"),
                ],
            },
            {
                "selector": "tbody th",
                "props": [
                    ("background-color", "#f7fbf7"),
                    ("color", "#6c7f70"),
                    ("border", f"1px solid {BORDER_COLOR}"),
                    ("padding", "10px 10px"),
                    ("font-weight", "600"),
                ],
            },
            {
                "selector": "td",
                "props": [
                    ("background-color", CARD_BG),
                    ("color", TEXT_COLOR),
                    ("border", f"1px solid {BORDER_COLOR}"),
                    ("padding", "10px 12px"),
                    ("font-size", "0.93rem"),
                ],
            },
            {
                "selector": "tbody tr:nth-child(even) td",
                "props": [("background-color", "#f8fcf7")],
            },
            {
                "selector": "tbody tr:nth-child(odd) td",
                "props": [("background-color", "#ffffff")],
            },
            {
                "selector": "tbody tr:hover td",
                "props": [("background-color", "#eef7ee")],
            },
            {
                "selector": "caption",
                "props": [
                    ("caption-side", "top"),
                    ("text-align", "left"),
                    ("font-weight", "600"),
                    ("color", "#48624d"),
                    ("padding", "0 0 8px 0"),
                ],
            },
        ]
    )


def style_score_cell(value):
    text_color = TEXT_COLOR if 0.4 <= float(value) < 0.7 else "white"
    return (
        f"background: linear-gradient(90deg, {color_scale(float(value))} 0%, {color_scale(float(value))}dd 100%); "
        f"color: {text_color}; font-weight: 800; border-left: 4px solid {color_scale(float(value))};"
    )


def style_breeding_score_cell(value):
    return (
        "background: linear-gradient(90deg, rgba(46, 125, 50, 0.08) 0%, rgba(46, 125, 50, 0.18) 100%); "
        f"color: {TEXT_COLOR}; font-weight: 800; border-left: 4px solid {LOW_STRESS_COLOR};"
    )


def style_genotype_cell(value, best=None, worst=None):
    if best is not None and value == best:
        return (
            "background-color: #edf7ed; "
            f"color: {LOW_STRESS_COLOR}; font-weight: 700; "
            f"border-left: 4px solid {LOW_STRESS_COLOR};"
        )
    if worst is not None and value == worst:
        return (
            "background-color: #fdeeee; "
            f"color: {HIGH_STRESS_COLOR}; font-weight: 700; "
            f"border-left: 4px solid {HIGH_STRESS_COLOR};"
        )
    return ""


def style_rank_cell(value):
    return (
        "background-color: rgba(46, 125, 50, 0.08); "
        f"color: {LOW_STRESS_COLOR}; font-weight: 800; text-align: center;"
    )


def style_class_cell(value):
    label = str(value).lower()
    if "low" in label:
        color = LOW_STRESS_COLOR
        text = "white"
    elif "moderate" in label:
        color = MODERATE_STRESS_COLOR
        text = TEXT_COLOR
    else:
        color = HIGH_STRESS_COLOR
        text = "white"
    return (
        f"background-color: {color}; color: {text}; font-weight: 700; "
        f"border-left: 4px solid {color};"
    )


def render_styled_table(
    df_table,
    stress_col=None,
    score_col=None,
    best=None,
    worst=None,
    caption_text=None,
):
    styler = style_dataframe(df_table.style)
    if stress_col and stress_col in df_table.columns:
        styler = styler.map(style_score_cell, subset=[stress_col]).format(
            {stress_col: "{:.3f}"}
        )
    if score_col and score_col in df_table.columns:
        styler = styler.map(style_breeding_score_cell, subset=[score_col]).format(
            {score_col: "{:.3f}"}
        )
    if "Genotype" in df_table.columns:
        styler = styler.map(
            lambda x: style_genotype_cell(x, best=best, worst=worst), subset=["Genotype"]
        )
    if "Rank" in df_table.columns:
        styler = styler.map(style_rank_cell, subset=["Rank"])
    if "Class" in df_table.columns:
        styler = styler.map(style_class_cell, subset=["Class"])
    if "Yield_Class" in df_table.columns:
        styler = styler.map(style_class_cell, subset=["Yield_Class"])
    if caption_text:
        st.caption(caption_text)
    st.dataframe(styler, use_container_width=True, hide_index=True)


def start_section(kicker):
    st.markdown(
        f"""
<div class="section-card">
  <div class="section-kicker">{kicker}</div>
""",
        unsafe_allow_html=True,
    )


def end_section():
    st.markdown("</div>", unsafe_allow_html=True)


def render_summary_tile(label, value, detail):
    st.markdown(
        f"""
<div class="summary-tile">
  <div class="summary-kicker">{label}</div>
  <div class="summary-value">{value}</div>
  <div class="summary-detail">{detail}</div>
</div>
""",
        unsafe_allow_html=True,
    )


def render_spotlight_card(selected, highlight, has_yield):
    stress_value = highlight["Stress_Index"].mean()
    stress_state = (
        "Low stress"
        if stress_value < 0.4
        else "Moderate stress" if stress_value < 0.7 else "High stress"
    )
    if has_yield and highlight["Yield"].notna().any():
        yield_text = f"{highlight['Yield'].mean():.2f}"
        yield_sentence = f"Average yield for this genotype is <b>{yield_text}</b>."
    else:
        yield_sentence = "Yield data is not available for this genotype in the current dataset."

    st.markdown(
        f"""
<div class="spotlight-card">
  <div class="spotlight-header">🔎 Genotype Spotlight</div>
  <div class="spotlight-value">{selected}</div>
  <div class="spotlight-note">
    Stress profile: <b>{stress_state}</b> with mean Stress Index <b>{stress_value:.3f}</b>.<br>
    {yield_sentence}
  </div>
</div>
""",
        unsafe_allow_html=True,
    )


def get_priority_columns(df_table):
    columns = ["Genotype", FVFM_COL, SPAD_COL]
    optional = [
        "Yield",
        "Stress_Index",
        "Breeding_Score",
        "Yield_Class",
        "Class",
    ]
    for col in optional:
        if col in df_table.columns:
            columns.append(col)
    return [col for col in columns if col in df_table.columns]


def get_stress_overview_columns(df_table):
    columns = ["Genotype", SPAD_COL, FVFM_COL, "Stress_Index", "Class"]
    if "Yield" in df_table.columns:
        columns.insert(3, "Yield")
    if "Breeding_Score" in df_table.columns:
        columns.append("Breeding_Score")
    return [col for col in columns if col in df_table.columns]


def render_leaderboard_cards(ranking_table):
    leaders = ranking_table.head(3)
    st.markdown('<div class="leader-grid">', unsafe_allow_html=True)
    cols = st.columns(max(len(leaders), 1))
    badges = ["🥇 Rank 1", "🥈 Rank 2", "🥉 Rank 3"]
    for idx, row in enumerate(leaders.itertuples(index=False)):
        with cols[idx]:
            st.markdown(
                f"""
<div class="leader-card">
  <div class="leader-rank">{badges[idx]}</div>
  <div class="leader-genotype">{row.Genotype}</div>
  <div class="leader-score">Stress Index: <b>{row.Stress_Index:.3f}</b></div>
</div>
""",
                unsafe_allow_html=True,
            )
    st.markdown("</div>", unsafe_allow_html=True)


def render_executive_card(title, body):
    st.markdown(
        f"""
<div class="executive-card">
  <div class="executive-title">{title}</div>
  <div class="executive-text">{body}</div>
</div>
""",
        unsafe_allow_html=True,
    )


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def format_report_heading(paragraph):
    if paragraph.runs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(31, 92, 43)
            run.font.bold = True


def format_report_caption(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if paragraph.runs:
        for run in paragraph.runs:
            run.italic = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(72, 98, 77)


def add_stress_legend(ax, include_selected=False):
    handles = [
        Line2D(
            [0],
            [0],
            marker="o",
            color="w",
            label="Low stress",
            markerfacecolor=LOW_STRESS_COLOR,
            markeredgecolor="black",
            markersize=8,
        ),
        Line2D(
            [0],
            [0],
            marker="o",
            color="w",
            label="Moderate stress",
            markerfacecolor=MODERATE_STRESS_COLOR,
            markeredgecolor="black",
            markersize=8,
        ),
        Line2D(
            [0],
            [0],
            marker="o",
            color="w",
            label="High stress",
            markerfacecolor=HIGH_STRESS_COLOR,
            markeredgecolor="black",
            markersize=8,
        ),
    ]
    if include_selected:
        handles.append(
            Line2D(
                [0],
                [0],
                marker="o",
                color="w",
                label="Selected genotype",
                markerfacecolor="blue",
                markeredgecolor="black",
                markersize=9,
            )
        )
    ax.legend(handles=handles, title="Stress class", frameon=False, loc="best")


def plot_stress(df):
    stress_df = df.groupby("Genotype", as_index=False)["Stress_Index"].mean()
    fig, ax = plt.subplots()

    colors = stress_df["Stress_Index"].apply(
        lambda x: (
            LOW_STRESS_COLOR
            if x < 0.4
            else MODERATE_STRESS_COLOR if x < 0.7 else HIGH_STRESS_COLOR
        )
    )

    bars = ax.bar(stress_df["Genotype"], stress_df["Stress_Index"], color=colors)
    for bar in bars:
        height = bar.get_height()
        ax.text(
            bar.get_x() + bar.get_width() / 2.0,
            height,
            f"{height:.2f}",
            ha="center",
            va="bottom",
        )

    ax.set_title("Stress Classification by Genotype")
    ax.set_xlabel("Genotype")
    ax.set_ylabel("Stress Index")
    add_stress_legend(ax)
    style_matplotlib(ax)
    return fig


def plot_yield(df):
    yield_vals = df.groupby("Genotype")["Yield"].mean()
    fig, ax = plt.subplots()
    bars = ax.bar(yield_vals.index, yield_vals.values, color=YIELD_COLOR)
    ax.set_title("Yield Performance by Genotype")
    ax.set_xlabel("Genotype")
    ax.set_ylabel("Yield (relative units)")

    for bar in bars:
        height = bar.get_height()
        ax.text(
            bar.get_x() + bar.get_width() / 2.0,
            height,
            f"{height:.2f}",
            ha="center",
            va="bottom",
        )

    style_matplotlib(ax)
    return fig


def plot_stress_vs_yield(df):
    from matplotlib.patches import Patch

    fig, ax = plt.subplots()

    colors = [
        get_color(
            df["Stress_Index"].iloc[i],
            df["Yield"].iloc[i],
            df,
        )
        for i in range(len(df))
    ]

    ax.scatter(df["Stress_Index"], df["Yield"], c=colors)

    for _, row in df.iterrows():
        ax.text(row["Stress_Index"], row["Yield"], row["Genotype"])

    ax.set_xlabel("Stress Index")
    ax.set_ylabel("Yield")
    ax.set_title("Stress vs Yield Breeding Map")
    ax.legend(
        handles=[
            Patch(facecolor=LOW_STRESS_COLOR, edgecolor="none", label="Elite"),
            Patch(
                facecolor=MODERATE_STRESS_COLOR,
                edgecolor="none",
                label="Intermediate",
            ),
            Patch(facecolor=HIGH_STRESS_COLOR, edgecolor="none", label="Poor"),
        ],
        title="Breeding Class",
        frameon=False,
        loc="best",
    )
    style_matplotlib(ax)
    fig.savefig("tradeoff.png", dpi=300, bbox_inches="tight")
    return fig


def run_yield_analysis(df):
    yield_rank = df.groupby("Genotype")["Yield"].mean().sort_values(ascending=False)
    fig = plot_yield(df)

    st.subheader("🌾 Yield Ranking (Relative Performance)")
    st.markdown(
        """
Yield classification is based on within-experiment distribution (percentiles), not fixed thresholds.
"""
    )
    render_styled_table(
        yield_rank.reset_index().rename(columns={"Yield": "Yield"}),
        best=df.groupby("Genotype")["Stress_Index"].mean().sort_values().index[0],
        worst=df.groupby("Genotype")["Stress_Index"].mean().sort_values().index[-1],
        caption_text="Higher yield values indicate stronger relative productivity in this experiment.",
    )

    st.subheader("🌾 Yield Performance")
    st.markdown(
        """
### 🌾 What this means

This plot shows productivity differences among genotypes.

- Higher bars = higher productivity  
- Lower bars = reduced performance  

👉 Yield helps balance **stress tolerance vs productivity trade-offs**.
"""
    )
    st.pyplot(fig)

    best_yield = yield_rank.idxmax()
    worst_yield = yield_rank.idxmin()
    st.subheader("🌾 Yield Insight")
    st.success(f"Best yield genotype: {best_yield}")
    st.error(f"Lowest yield genotype: {worst_yield}")


def run_tradeoff(df):
    explain_stress_vs_yield()
    st.pyplot(plot_stress_vs_yield(df))

    st.subheader("🧠 Breeding Interpretation")

    elite = df[
        (df["Stress_Index"] < df["Stress_Index"].quantile(0.33))
        & (df["Yield"] > df["Yield"].quantile(0.66))
    ]
    sensitive = df[
        (df["Stress_Index"] > df["Stress_Index"].quantile(0.66))
        & (df["Yield"] < df["Yield"].quantile(0.33))
    ]

    st.success(f"🌱 Elite genotypes (recommended): {', '.join(elite['Genotype'].unique())}")
    st.error(f"🔥 Sensitive genotypes (avoid): {', '.join(sensitive['Genotype'].unique())}")


def generate_word_report(df):
    has_yield = safe_yield_available(df)
    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)

    ranking = df.groupby("Genotype")["Stress_Index"].mean().sort_values()
    best = ranking.idxmin()
    worst = ranking.idxmax()

    title = doc.add_paragraph()
    title_run = title.add_run("PhytoStress AI Scientific Report")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(31, 92, 43)

    intro = doc.add_paragraph()
    intro.add_run(
        "Technical summary of genotype performance under drought stress using "
    )
    intro.add_run("Chlorophyll Content, Photosystem II Efficiency, Stress Index").bold = True
    intro.add_run(" and ")
    intro.add_run("Yield").bold = True
    intro.add_run(" when available.")

    doc.add_paragraph(
        "This document preserves the analytical outputs of the dashboard and organizes them as a concise scientific report for interpretation and decision support."
    )

    heading = doc.add_heading("1. Genotype Ranking", level=1)
    format_report_heading(heading)
    doc.add_paragraph(
        """
This table shows the genotypes ranked by Stress Index.

- Lower values = better drought tolerance
- Higher values = more sensitive plants

The best performing genotype is the one with the lowest Stress Index.
"""
    )
    subheading = doc.add_heading("Genotype Ranking Table", level=2)
    format_report_heading(subheading)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Rank"
    hdr[1].text = "Genotype"
    hdr[2].text = "Stress Index"
    for cell in hdr:
        set_cell_shading(cell, "2E7D32")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    for i, (geno, val) in enumerate(ranking.items(), 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = str(geno)
        row[2].text = f"{val:.3f}"
        if geno == best:
            for cell in row:
                set_cell_shading(cell, "EDF7ED")
        elif geno == worst:
            for cell in row:
                set_cell_shading(cell, "FDEEEE")

    heading = doc.add_heading("2. Stress Index Explanation", level=1)
    format_report_heading(heading)
    doc.add_paragraph(
        """
The Stress Index summarizes drought response using normalized physiological traits.

It is calculated as:
Stress Index = 1 - (normalized Chlorophyll Content + normalized Photosystem II Efficiency) / 2

- Values close to 0 → more tolerant plants
- Values close to 1 → more stressed plants
"""
    )

    if "Yield" in df.columns:
        heading = doc.add_heading("3. Yield Performance", level=1)
        format_report_heading(heading)
        doc.add_paragraph(
            """
Yield represents the productivity potential of each genotype.

In this analysis, yield is interpreted relatively:
- High yield = better performance in this experiment
- Low yield = lower productivity under the same conditions
"""
        )

    heading = doc.add_heading("4. Heatmap Interpretation", level=1)
    format_report_heading(heading)
    doc.add_paragraph(
        """
The heatmap shows relationships between Stress Index, Chlorophyll Content, and Photosystem II Efficiency.

- Green colors indicate better plant performance
- Red colors indicate higher stress

This helps visualize overall genotype behavior in a single view.
"""
    )

    if "Yield" in df.columns:
        heading = doc.add_heading("5. Stress vs Yield Relationship", level=1)
        format_report_heading(heading)
        doc.add_paragraph(
            """
This graph compares stress tolerance and yield at the same time.

- Ideal genotypes appear with low stress and high yield
- Poor genotypes show high stress and low yield

This is the most important figure for selecting elite genotypes.
"""
        )

    fig1_path = "fig1.png"
    fig1.savefig(fig1_path, dpi=300, bbox_inches="tight")

    fig2_path = "fig2.png"
    fig2.savefig(fig2_path, dpi=300, bbox_inches="tight")

    heading = doc.add_heading("Figures", level=1)
    format_report_heading(heading)
    cap = doc.add_paragraph("Figure 1. Stress Index by Genotype.")
    format_report_caption(cap)
    doc.add_picture(fig1_path, width=Inches(5.5))

    cap = doc.add_paragraph("Figure 2. Chlorophyll Content versus Photosystem II Efficiency.")
    format_report_caption(cap)
    doc.add_picture(fig2_path, width=Inches(5.5))

    cap = doc.add_paragraph("Figure 3. Integrated heatmap of physiological traits and stress response.")
    format_report_caption(cap)
    doc.add_picture("heatmap.png", width=Inches(5.5))

    if has_yield:
        cap = doc.add_paragraph("Figure 4. Stress versus Yield trade-off for breeding decisions.")
        format_report_caption(cap)
        doc.add_picture("tradeoff.png", width=Inches(5.5))

    low = df[df["Stress_Index"] < 0.4]["Genotype"].unique()
    mid = df[
        (df["Stress_Index"] >= 0.4) & (df["Stress_Index"] < 0.7)
    ]["Genotype"].unique()
    high = df[df["Stress_Index"] >= 0.7]["Genotype"].unique()

    heading = doc.add_heading("Stress Classification", level=1)
    format_report_heading(heading)
    doc.add_paragraph("Low stress (tolerant): " + ", ".join(low))
    doc.add_paragraph("Moderate stress: " + ", ".join(mid))
    doc.add_paragraph("High stress (sensitive): " + ", ".join(high))

    heading = doc.add_heading("Discussion", level=1)
    format_report_heading(heading)
    if "Breeding_Score" in df.columns and has_yield:
        final_rank = df.groupby("Genotype")["Breeding_Score"].mean().sort_values(
            ascending=False
        )
        report_best = final_rank.idxmax()
        report_worst = final_rank.idxmin()
        avg_stress = df["Stress_Index"].mean()

        if avg_stress < 0.4:
            stress_state = "low overall stress conditions"
        elif avg_stress < 0.7:
            stress_state = "moderate stress conditions"
        else:
            stress_state = "high stress conditions"

        doc.add_paragraph(
            f"""
The evaluated genotypes were analyzed under {stress_state}, showing variability in physiological and agronomic performance.

The breeding score analysis identified {report_best} as the most promising genotype, while {report_worst} showed the lowest performance.

These results confirm that integrating Stress Index and Yield is effective for genotype selection under drought conditions.
"""
        )
    else:
        doc.add_paragraph(
            f"""
The evaluated genotypes were analyzed for physiological stress response under drought conditions.

The ranking analysis identified {best} as the most promising genotype, while {worst} showed the lowest performance.

These results confirm that physiological indicators such as Chlorophyll Content, Photosystem II Efficiency, and Stress Index are effective for genotype selection under drought conditions.
"""
        )

    heading = doc.add_heading("Final Recommendation", level=1)
    format_report_heading(heading)
    if "Breeding_Score" in df.columns and has_yield:
        final_rank = df.groupby("Genotype")["Breeding_Score"].mean().sort_values(
            ascending=False
        )
        report_best = final_rank.idxmax()
        report_worst = final_rank.idxmin()
        paragraph = doc.add_paragraph()
        paragraph.add_run("Best genotype: ").bold = True
        paragraph.add_run(f"{report_best}\n")
        paragraph.add_run("Worst genotype: ").bold = True
        paragraph.add_run(f"{report_worst}\n\n")
        paragraph.add_run("Best genotype is recommended because it combines:\n")
        paragraph.add_run("• Low Stress Index\n• High Yield\n• Strong overall Breeding Score\n\n")
        paragraph.add_run(
            "This genotype represents the ideal candidate for drought tolerance selection."
        )
    else:
        paragraph = doc.add_paragraph()
        paragraph.add_run("Best genotype: ").bold = True
        paragraph.add_run(f"{best}\n")
        paragraph.add_run("Worst genotype: ").bold = True
        paragraph.add_run(f"{worst}\n\n")
        paragraph.add_run("Best genotype is recommended because it combines:\n")
        paragraph.add_run("• Low stress response\n• High physiological stability\n• Strong drought tolerance\n\n")
        paragraph.add_run(
            "This genotype represents the ideal candidate for drought tolerance selection."
        )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_full_report(df, ranking):
    doc = Document()

    title = doc.add_heading("PhytoStress AI Full Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Drought stress phenotyping integrated report")
    doc.add_page_break()

    doc.add_heading("1. Genotype Ranking", level=1)
    for i, (g, v) in enumerate(ranking.items(), 1):
        doc.add_paragraph(f"{i}. {g}: {v:.3f}")

    fig1_path = "fig1.png"
    fig2_path = "fig2.png"
    fig3_path = "heatmap.png"

    doc.add_heading("2. Figures", level=1)
    doc.add_picture(fig1_path, width=Inches(5.5))
    doc.add_picture(fig2_path, width=Inches(5.5))
    doc.add_picture(fig3_path, width=Inches(5.5))

    doc.add_heading("3. Stress Classification Summary", level=1)

    low = df[df["Stress_Index"] < 0.4]["Genotype"].unique()
    mid = df[(df["Stress_Index"] >= 0.4) & (df["Stress_Index"] < 0.7)]["Genotype"].unique()
    high = df[df["Stress_Index"] >= 0.7]["Genotype"].unique()

    doc.add_paragraph("Low stress: " + ", ".join(low))
    doc.add_paragraph("Moderate stress: " + ", ".join(mid))
    doc.add_paragraph("High stress: " + ", ".join(high))

    best = ranking.idxmin()
    worst = ranking.idxmax()

    doc.add_heading("4. Conclusion", level=1)
    doc.add_paragraph(
        f"The analysis identified {best} as the most drought-tolerant genotype "
        f"and {worst} as the most sensitive genotype. "
        "Physiological traits (Chlorophyll Content and Photosystem II Efficiency) successfully differentiated stress responses."
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


# =========================
# MAIN APP FLOW
# =========================
col1, col2 = st.columns(2)

with col1:
    uploaded_file = st.file_uploader("📂 Upload Excel file (.xlsx)", type=["xlsx"])

with col2:
    run_demo = st.button("🚀 Run Demo")

df = None

if run_demo:
    df = pd.DataFrame(
        {
            "Genotype": ["H1", "H2", "H3", "H4"],
            "FvFm": [0.80, 0.72, 0.58, 0.61],
            "SPAD": [45, 40, 30, 33],
            "Yield": [6.1, 5.4, 4.0, 4.2],
        }
    )
    st.success("Demo loaded successfully")

elif uploaded_file is not None:
    df = pd.read_excel(uploaded_file)
    st.success("File uploaded successfully")
    render_styled_table(df[get_priority_columns(df)])
else:
    st.stop()

df = normalize_trait_columns(df)

required = ["Genotype", FVFM_COL, SPAD_COL]

if any(col not in df.columns for col in required):
    st.error(
        f"Missing required columns: Genotype, {FVFM_COL}, {SPAD_COL}"
    )
    st.stop()

df = df.dropna(subset=["Genotype", FVFM_COL, SPAD_COL])
df["Genotype"] = df["Genotype"].astype(str)

df = compute_stress(df)

if uploaded_file or run_demo:
    st.subheader("📂 Simulated Dataset" if run_demo else "📋 Data Preview")
    render_styled_table(
        df[get_priority_columns(df)],
        caption_text="Compact data preview with the core variables used in the analysis pipeline.",
    )

if "Yield" in df.columns:
    df["Yield"] = pd.to_numeric(df["Yield"], errors="coerce")
    df = df.dropna(subset=["Yield"])
else:
    df["Yield"] = np.nan

for col in df.columns:
    if "yield" in col.lower() and "estimated" in col.lower():
        df[col] = np.nan

has_yield = "Yield" in df.columns and df["Yield"].notna().any()

if has_yield:
    start_section("03 • Yield & Trade-off")
    q33 = df["Yield"].quantile(0.33)
    q66 = df["Yield"].quantile(0.66)

    def yield_class(x):
        if x >= q66:
            return "🟢 High Yield"
        if x >= q33:
            return "🟡 Medium Yield"
        return "🔴 Low Yield"

    def yield_color(x):
        if x >= q66:
            return LOW_STRESS_COLOR
        if x >= q33:
            return MODERATE_STRESS_COLOR
        return HIGH_STRESS_COLOR

    df["Yield_Class"] = df["Yield"].apply(yield_class)
    df["Stress_norm"] = 1 - (
        (df["Stress_Index"] - df["Stress_Index"].min())
        / (df["Stress_Index"].max() - df["Stress_Index"].min())
    )
    df["Yield_norm"] = (
        (df["Yield"] - df["Yield"].min())
        / (df["Yield"].max() - df["Yield"].min())
    )
    df["Breeding_Score"] = 0.5 * df["Stress_norm"] + 0.5 * df["Yield_norm"]

# =========================
# RANKING
# =========================
st.subheader("📊 Genotype Display Control")

top_n = st.selectbox(
    "Show top genotypes",
    options=[10, 15, 20, 30, "All"],
    index=0,
)

ranking = df.groupby("Genotype")["Stress_Index"].mean().sort_values()
ranking_df = ranking.reset_index()
ranking_df.columns = ["Genotype", "Stress_Index"]
ranking_df = ranking_df.sort_values("Stress_Index")
best = ranking.idxmin()
worst = ranking.idxmax()
avg_stress = df["Stress_Index"].mean()
top_yield_text = (
    df.groupby("Genotype")["Yield"].mean().idxmax() if has_yield else "N/A"
)
fig1 = plot_stress(df)
fig2 = plot_yield(df) if has_yield else None
fig_tradeoff = plot_stress_vs_yield(df) if has_yield else None

summary_col, spotlight_col = st.columns([2.2, 1], gap="large")

with summary_col:
    st.markdown('<div class="metrics-card">', unsafe_allow_html=True)
    st.subheader("📈 Dashboard Summary")
    st.markdown(
        '<div class="metrics-note">Key result snapshot for fast screening before exploring the detailed tables and plots.</div>',
        unsafe_allow_html=True,
    )
    tile_cols = st.columns(4)
    with tile_cols[0]:
        render_summary_tile(
            "🌱 Best Genotype",
            best,
            "Lowest mean Stress Index across the evaluated genotypes.",
        )
    with tile_cols[1]:
        render_summary_tile(
            "🔥 Worst Genotype",
            worst,
            "Highest mean Stress Index and weakest drought stability.",
        )
    with tile_cols[2]:
        render_summary_tile(
            "🧪 Average Stress",
            f"{avg_stress:.3f}",
            "Dataset-level physiological stress summary.",
        )
    with tile_cols[3]:
        render_summary_tile(
            "🌾 Top Yield",
            top_yield_text,
            "Highest mean productivity observed in the current dataset.",
        )
    st.markdown("</div>", unsafe_allow_html=True)

with spotlight_col:
    st.subheader("🔎 Search Genotype")
    genotype_list = df["Genotype"].unique()
    selected = st.selectbox(
        "Find specific genotype",
        genotype_list,
    )
    highlight = df[df["Genotype"] == selected]
    render_spotlight_card(selected, highlight, has_yield)


start_section("01 • Ranking & Selection")

def medal(i):
    if i == 0:
        return "🥇"
    if i == 1:
        return "🥈"
    if i == 2:
        return "🥉"
    return ""

st.subheader("🏆 Genotype Ranking")
for i, row in enumerate(ranking_df.itertuples()):
    color = color_scale(row.Stress_Index)
    st.markdown(
        f"<div style='padding:10px 12px;background-color:{color}18;border:1px solid {color}; border-radius:12px; margin-bottom:6px;'>"
        f"{medal(i)} <b>{row.Genotype}</b> → {row.Stress_Index:.3f}"
        f"</div>",
        unsafe_allow_html=True,
    )

ranking_table = ranking_df.copy()
ranking_table.insert(0, "Rank", range(1, len(ranking_table) + 1))
render_leaderboard_cards(ranking_df)
render_styled_table(
    ranking_table,
    stress_col="Stress_Index",
    best=best,
    worst=worst,
    caption_text="Lower Stress Index values indicate stronger drought tolerance. The top three cards provide the fastest shortlist for breeding selection.",
)

st.subheader("🥇 Breeding Recommendation")

st.markdown(
    """
### 🧠 How to read this section

This section identifies the best genotypes for breeding based on the lowest Stress Index.

- Top ranked genotypes → most stable under drought stress  
- Bottom ranked genotypes → poor performance under stress  

👉 These are the **final candidates for selection in breeding programs**.
"""
)

top3 = ranking.nsmallest(3)

for i, (g, v) in enumerate(top3.items(), 1):
    st.markdown(f"**{i}. {g} → Stress Index: {v:.3f}**")

st.success(
    f"""
🌱 Key Message:
Genotype {best} shows superior drought tolerance and is the strongest candidate for breeding improvement.

🔥 Genotype {worst} is highly sensitive and not recommended for drought-prone environments.
"""
)

end_section()

if "Breeding_Score" in df.columns:
    start_section("02 • Breeding Score")
    st.subheader("🏆 Final Breeding Score Ranking")
    final_rank = df.groupby("Genotype")["Breeding_Score"].mean().sort_values(ascending=False)
    render_styled_table(
        final_rank.reset_index().rename(columns={"Breeding_Score": "Breeding_Score"}),
        score_col="Breeding_Score",
        caption_text="Higher breeding scores indicate better overall balance between drought tolerance and productivity.",
    )
    best_genotype = final_rank.idxmax()
    st.success(
        f"""
🌱 Best overall breeding genotype: {best_genotype}

This genotype shows the best combination of:
- Low drought stress
- High yield performance

It is recommended as the primary candidate for breeding programs.
"""
    )

    st.subheader("🌾 Stress vs Yield with Breeding Classification")

    fig_breed, ax_breed = plt.subplots()
    elite_cut = df["Breeding_Score"].quantile(0.66)
    poor_cut = df["Breeding_Score"].quantile(0.33)
    breed_colors = []
    breed_labels = []

    for _, row in df.iterrows():
        if row["Breeding_Score"] >= elite_cut:
            label = "🟢 Elite"
            color = LOW_STRESS_COLOR
        elif row["Breeding_Score"] >= poor_cut:
            label = "🟡 Good"
            color = MODERATE_STRESS_COLOR
        else:
            label = "🔴 Poor"
            color = HIGH_STRESS_COLOR

        breed_labels.append(label)
        breed_colors.append(color)

    ax_breed.scatter(
        df["Stress_Index"],
        df["Yield"],
        c=breed_colors,
        s=120,
        alpha=0.85,
        edgecolors="black",
        linewidths=0.6,
    )

    for (_, row), label, color in zip(df.iterrows(), breed_labels, breed_colors):
        ax_breed.text(
            row["Stress_Index"],
            row["Yield"],
            f"{row['Genotype']}\n{label}",
            fontsize=8,
            color=color,
        )

    ax_breed.set_xlabel("Stress Index (Lower = Better)")
    ax_breed.set_ylabel("Yield (Higher = Better)")
    ax_breed.set_title("Genotype Breeding Space")

    style_matplotlib(ax_breed)
    ax_breed.legend(
        handles=[
            plt.Line2D(
                [0],
                [0],
                marker="o",
                color="w",
                label="Elite",
                markerfacecolor=LOW_STRESS_COLOR,
                markeredgecolor="black",
                markersize=9,
            ),
            plt.Line2D(
                [0],
                [0],
                marker="o",
                color="w",
                label="Good",
                markerfacecolor=MODERATE_STRESS_COLOR,
                markeredgecolor="black",
                markersize=9,
            ),
            plt.Line2D(
                [0],
                [0],
                marker="o",
                color="w",
                label="Poor",
                markerfacecolor=HIGH_STRESS_COLOR,
                markeredgecolor="black",
                markersize=9,
            ),
        ],
        title="Breeding Class",
        frameon=False,
        loc="best",
    )
    st.pyplot(fig_breed)
    end_section()

if has_yield:
    st.subheader("🌾 Yield Analysis")
    run_yield_analysis(df)
    run_tradeoff(df)
    end_section()
else:
    st.warning("Yield data not available. Yield analysis is disabled.")

if "Breeding_Score" in df.columns and has_yield:
    st.subheader("📄 Automated Discussion (Scientific Interpretation)")

    best = final_rank.idxmax()
    worst = final_rank.idxmin()

    avg_stress = df["Stress_Index"].mean()
    avg_yield = df["Yield"].mean()

    if avg_stress < 0.4:
        stress_state = "low overall stress conditions"
    elif avg_stress < 0.7:
        stress_state = "moderate stress conditions"
    else:
        stress_state = "high stress conditions"

    if avg_yield > df["Yield"].median():
        yield_state = "above-average yield performance"
    else:
        yield_state = "below-average yield performance"

    st.markdown(
        f"""
### Discussion

The evaluated genotypes were analyzed under {stress_state}, showing variability in both physiological stress response and yield potential.

Overall, the dataset indicates {yield_state}, suggesting that environmental conditions and genotype interactions significantly influenced performance.

The breeding score analysis identified **{best}** as the most promising genotype due to its superior combination of low stress and high yield.

Conversely, **{worst}** showed the lowest performance and is not recommended for drought-prone environments.

These results support the use of integrated physiological and agronomic indicators (Chlorophyll Content, Photosystem II Efficiency, Stress Index, and Yield) for efficient selection of drought-tolerant genotypes.
"""
    )

# =========================
# CLASSIFICATION
# =========================
start_section("04 • Stress Classification")
df["Class"] = df["Stress_Index"].apply(classify)
low = df[df["Stress_Index"] < 0.4]
mid = df[(df["Stress_Index"] >= 0.4) & (df["Stress_Index"] < 0.7)]
high = df[df["Stress_Index"] >= 0.7]

st.subheader("📊 Stress Classification")
st.markdown(
    """
### 🧠 What this means

Genotypes are grouped based on Stress Index values:

- 🟢 Low stress → drought tolerant (preferred for breeding)
- 🟡 Moderate stress → intermediate response
- 🔴 High stress → sensitive genotypes (not recommended)

👉 Each point represents one genotype classified by stress level.
"""
)
render_styled_table(
    df[get_stress_overview_columns(df)],
    stress_col="Stress_Index",
    best=best,
    worst=worst,
    caption_text="Stress Index values are the main drought-tolerance indicator: low = tolerant, moderate = intermediate, high = sensitive.",
)

tab1, tab2, tab3 = st.tabs(["🟢 Low", "🟡 Moderate", "🔴 High"])

with tab1:
    render_styled_table(
        low[get_stress_overview_columns(low)],
        stress_col="Stress_Index",
        best=best,
        worst=worst,
        caption_text="Low-stress genotypes are the strongest candidates for drought-focused breeding.",
    )

with tab2:
    render_styled_table(
        mid[get_stress_overview_columns(mid)],
        stress_col="Stress_Index",
        best=best,
        worst=worst,
        caption_text="Moderate-stress genotypes show intermediate physiological stability.",
    )

with tab3:
    render_styled_table(
        high[get_stress_overview_columns(high)],
        stress_col="Stress_Index",
        best=best,
        worst=worst,
        caption_text="High-stress genotypes are more sensitive and lower priority for drought selection.",
    )

st.subheader("📊 Stress Classes Overview")

fig, ax = plt.subplots()

colors = df["Stress_Index"].apply(
    lambda x: LOW_STRESS_COLOR if x < 0.4 else MODERATE_STRESS_COLOR if x < 0.7 else HIGH_STRESS_COLOR
)

ax.scatter(df["Genotype"], df["Stress_Index"], c=colors)

ax.set_title("Stress Classification by Genotype")
ax.set_xlabel("Genotype")
ax.set_ylabel("Stress Index")

for i in range(len(df)):
    ax.text(
        df["Genotype"].iloc[i],
        df["Stress_Index"].iloc[i],
        df["Genotype"].iloc[i],
    )

add_stress_legend(ax)
style_matplotlib(ax)
st.pyplot(fig)

end_section()

start_section("05 • Integrated Heatmap")

st.subheader("🌡️ Integrated Stress Heatmap")
st.markdown(
    """
### 🔬 What this heatmap shows

This heatmap combines all physiological traits and stress response.

- 🟢 Green → optimal performance  
- 🟡 Yellow → intermediate response  
- 🔴 Red → high stress sensitivity  

👉 It provides a **quick visual summary of genotype performance**.
"""
)

st.subheader("🌡️ Stress Heatmap (Genotype × Traits)")
heat = df.groupby("Genotype")[[FVFM_COL, SPAD_COL, "Stress_Index"]].mean()

fig, ax = plt.subplots()
heatmap = sns.heatmap(heat, annot=True, cmap="RdYlGn_r", ax=ax)
style_matplotlib(ax)
if heatmap.collections:
    style_colorbar(heatmap.collections[0].colorbar)
fig.savefig("heatmap.png", dpi=300, bbox_inches="tight")
st.pyplot(fig)
end_section()

# =========================
# RECOMMENDATION
# =========================
st.subheader("💡 Recommendation")
start_section("06 • Scientific Summary")
render_executive_card(
    "Recommendation",
    generate_recommendation(df["Stress_Index"].mean()),
)

st.subheader("🧠 Scientific Summary")
st.markdown(
    """
### What this means biologically

This section summarizes the biological meaning of the results and supports the final breeding decision.

- Identifies the most tolerant genotype
- Identifies the most sensitive genotype
- Summarizes how the physiological traits support stress discrimination
"""
)

avg = df["Stress_Index"].mean()

if avg < 0.4:
    stress_level = "low overall stress conditions"
elif avg < 0.7:
    stress_level = "moderate stress conditions"
else:
    stress_level = "high stress conditions"

render_executive_card(
    "Key findings",
    f"<b>🌱 Best genotype:</b> {best}<br>"
    f"<b>🔥 Worst genotype:</b> {worst}<br>"
    f"<b>📊 Dataset stress level:</b> {stress_level}",
)

render_executive_card(
    "Scientific interpretation",
    f"The analysis indicates <b>{stress_level}</b> based on the combined physiological response of "
    f"Chlorophyll Content and Photosystem II Efficiency.<br><br>"
    f"<b>Most tolerant genotype:</b> {best}<br>"
    f"<b>Most sensitive genotype:</b> {worst}<br><br>"
    "Lower Stress Index values indicate genotypes that maintained stronger physiological performance. "
    "These results support the use of integrated physiological indicators for selecting drought-tolerant breeding material.",
)

end_section()

# =========================
# BAR PLOT (GREEN - como ayer)
# =========================
st.subheader("📉 Stress Index by Genotype")

st.markdown(
    """
### 📊 What this graph shows

This plot ranks each genotype based on its Stress Index.

- 🟢 Low values = drought tolerant genotypes  
- 🟡 Medium values = intermediate response  
- 🔴 High values = drought sensitive genotypes  

👉 This is the **main selection criterion** for breeding decisions.
"""
)

start_section("07 • Visualization")
st.pyplot(fig1)

# =========================
# SCATTER (AZUL - como ayer)
# =========================
st.subheader("🌿 Physiological Relationship: Chlorophyll Content vs Photosystem Efficiency")
st.markdown(
    """
### 🌱 Biological meaning

This graph shows plant physiological health.

- Chlorophyll Content → leaf greenness and chlorophyll concentration  
- Photosystem II Efficiency → photosynthetic efficiency  

👉 Genotypes in the upper range indicate **healthier plants under stress**.

### 🧠 Trait definition update

- **Chlorophyll Content:** proxy of leaf greenness and chlorophyll level  
- **Photosystem II Efficiency:** indicator of photosynthetic performance and stress damage  
- **Yield:** agronomic output used to evaluate whether physiological tolerance is expressed as real productivity  

👉 These indicators are complementary but NOT equivalent.
"""
)

st.subheader("🌿 Physiological Trait Relationship")

fig2, ax2 = plt.subplots()

if top_n == "All":
    df_plot = df.copy()
else:
    df_plot = df[
        df["Genotype"].isin(
            df.groupby("Genotype")["Stress_Index"]
            .mean()
            .nsmallest(int(top_n))
            .index
        )
    ]

ax2.scatter(
    df_plot[SPAD_COL],
    df_plot[FVFM_COL],
    c=df_plot["Stress_Index"].apply(color_scale),
    edgecolors="black",
    linewidths=0.4,
)

ax2.scatter(
    highlight[SPAD_COL],
    highlight[FVFM_COL],
    color="blue",
    s=200,
    edgecolor="black",
)

ax2.set_xlabel("Chlorophyll Content")
ax2.set_ylabel("Photosystem II Efficiency")
add_stress_legend(ax2, include_selected=True)
style_matplotlib(ax2)

st.pyplot(fig2)
end_section()

# =========================
# WORD REPORT BUTTON
# =========================
if st.button("📄 Generate Full Scientific Report"):
    buffer = generate_word_report(df)
    st.download_button(
        "Download Word Report",
        buffer,
        file_name="PhytoStress_Report.docx",
    )

if st.button("Generate full report (disabled)"):
    pass

st.markdown(
    """
<div class="feature-strip">
  <div class="feature-card">
    <strong>🧪 Science-based</strong>
    <span>Integrated physiological and agronomic screening for drought evaluation.</span>
  </div>
  <div class="feature-card">
    <strong>🌾 Crop-independent</strong>
    <span>Suitable for any species with Chlorophyll Content, Photosystem II Efficiency, and Yield data.</span>
  </div>
  <div class="feature-card">
    <strong>🧠 AI-assisted insights</strong>
    <span>Turns trait measurements into interpretable breeding-oriented conclusions.</span>
  </div>
  <div class="feature-card">
    <strong>🎯 Selection support</strong>
    <span>Highlights elite genotypes and clarifies stress-versus-productivity trade-offs.</span>
  </div>
</div>
""",
    unsafe_allow_html=True,
)
